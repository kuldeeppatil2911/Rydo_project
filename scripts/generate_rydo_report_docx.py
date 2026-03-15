from pathlib import Path
import re
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parents[1]
SOURCE_MD = BASE_DIR / "Rydo_Project_Report_Reference_Style.md"
ASSET_DIR = BASE_DIR / "report_assets"
DOCX_PATH = BASE_DIR / "Rydo_Project_Report_Submission.docx"

TITLE = "RYDO - RIDE BOOKING AND SAFETY ALERT SYSTEM"
INSTITUTE = "Parul Institute of Technology"
DEPARTMENT = "Department of Computer Science and Engineering (CSE Core)"
STUDENT = "Kuldeep Ravindra Patil"
MENTOR = "Mr. Tadikonda Venkatata Durga Prasad"
ACADEMIC_YEAR = "2025-2026"


BLUE_DARK = (9, 31, 74)
BLUE = (26, 91, 164)
CYAN = (59, 180, 219)
LIME = (176, 216, 34)
WHITE = (255, 255, 255)
BLACK = (0, 0, 0)
TEXT = (43, 54, 66)
MUTED = (94, 110, 126)
CARD = (252, 253, 255)
LINE = (189, 202, 219)
NAVY = (19, 45, 79)
TEAL = (22, 146, 153)
ORANGE = (242, 139, 65)
PURPLE = (111, 95, 208)
GREEN = (45, 161, 110)


def load_font(size, bold=False):
    candidates = [
        "arialbd.ttf" if bold else "arial.ttf",
        "calibrib.ttf" if bold else "calibri.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


FONT_54_B = load_font(54, True)
FONT_34_B = load_font(34, True)
FONT_28_B = load_font(28, True)
FONT_24_B = load_font(24, True)
FONT_22 = load_font(22, False)
FONT_22_B = load_font(22, True)
FONT_18 = load_font(18, False)
FONT_18_B = load_font(18, True)
FONT_16 = load_font(16, False)
FONT_16_B = load_font(16, True)
FONT_14 = load_font(14, False)
FONT_14_B = load_font(14, True)


def wrap_lines(draw, text, font, max_width):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        trial = word if not current else f"{current} {word}"
        width = draw.textbbox((0, 0), trial, font=font)[2]
        if width <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [""]


def create_canvas(title, subtitle=""):
    img = Image.new("RGB", (1600, 900), (244, 247, 250))
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1600, 18), fill=BLACK)
    draw.rectangle((0, 18, 1600, 120), fill=BLUE)
    draw.rectangle((1440, 18, 1535, 120), fill=LIME)
    draw.rounded_rectangle((50, 42, 770, 92), radius=18, fill=CYAN)
    draw.text((70, 48), title, font=FONT_34_B, fill=WHITE)
    if subtitle:
        draw.text((70, 94), subtitle, font=FONT_16, fill=WHITE)
    draw.line((0, 122, 1600, 122), fill=CYAN, width=5)
    return img, draw


def draw_centered_text(draw, box, text, font, fill):
    x1, y1, x2, y2 = box
    lines = wrap_lines(draw, text, font, x2 - x1 - 20)
    heights = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        heights.append(bbox[3] - bbox[1])
    total_h = sum(heights) + (len(lines) - 1) * 6
    y = y1 + ((y2 - y1 - total_h) / 2)
    for line, height in zip(lines, heights):
        bbox = draw.textbbox((0, 0), line, font=font)
        width = bbox[2] - bbox[0]
        x = x1 + ((x2 - x1 - width) / 2)
        draw.text((x, y), line, font=font, fill=fill)
        y += height + 6


def box(draw, coords, text, fill, outline=None, font=FONT_18_B, text_fill=WHITE, radius=26):
    draw.rounded_rectangle(coords, radius=radius, fill=fill, outline=outline or fill, width=3)
    draw_centered_text(draw, coords, text, font, text_fill)


def outline_card(draw, coords, title, lines, accent):
    draw.rounded_rectangle(coords, radius=24, fill=CARD, outline=accent, width=3)
    x1, y1, x2, y2 = coords
    draw.text((x1 + 16, y1 + 12), title, font=FONT_18_B, fill=accent)
    y = y1 + 52
    for line in lines:
        wrapped = wrap_lines(draw, line, FONT_14, x2 - x1 - 38)
        for idx, part in enumerate(wrapped):
            prefix = u"\u2022 " if idx == 0 else "  "
            draw.text((x1 + 18, y), prefix + part, font=FONT_14, fill=TEXT)
            y += 24


def arrow(draw, p1, p2, color=LINE, width=5):
    draw.line((p1, p2), fill=color, width=width)
    x1, y1 = p1
    x2, y2 = p2
    if abs(x2 - x1) >= abs(y2 - y1):
        if x2 >= x1:
            pts = [(x2, y2), (x2 - 18, y2 - 10), (x2 - 18, y2 + 10)]
        else:
            pts = [(x2, y2), (x2 + 18, y2 - 10), (x2 + 18, y2 + 10)]
    else:
        if y2 >= y1:
            pts = [(x2, y2), (x2 - 10, y2 - 18), (x2 + 10, y2 - 18)]
        else:
            pts = [(x2, y2), (x2 - 10, y2 + 18), (x2 + 10, y2 + 18)]
    draw.polygon(pts, fill=color)


def actor(draw, x, y, color):
    draw.ellipse((x + 20, y, x + 70, y + 50), outline=color, width=4, fill=WHITE)
    draw.line((x + 45, y + 50, x + 45, y + 130), fill=color, width=4)
    draw.line((x, y + 78, x + 90, y + 78), fill=color, width=4)
    draw.line((x + 45, y + 130, x + 10, y + 175), fill=color, width=4)
    draw.line((x + 45, y + 130, x + 80, y + 175), fill=color, width=4)


def diagram_system_architecture(path):
    img, draw = create_canvas("System Architecture", "Frontend, backend, database, and external services")
    box(draw, (90, 235, 290, 325), "User", NAVY, font=FONT_22_B)
    box(draw, (360, 205, 670, 355), "React Frontend\nLogin | Dashboard | Book | Track | Profile", TEAL, font=FONT_22_B)
    box(draw, (760, 205, 1020, 355), "Express Backend\nREST API + JWT", ORANGE, font=FONT_24_B)
    box(draw, (1110, 150, 1400, 255), "MongoDB Atlas", GREEN, font=FONT_24_B)
    box(draw, (1110, 295, 1400, 400), "SMTP / Twilio", PURPLE, font=FONT_22_B)
    box(draw, (1110, 440, 1400, 545), "Vercel + Railway", NAVY, font=FONT_22_B)
    arrow(draw, (290, 280), (360, 280), TEAL)
    arrow(draw, (670, 280), (760, 280), TEAL)
    arrow(draw, (1020, 250), (1110, 200), TEAL)
    arrow(draw, (1020, 300), (1110, 345), TEAL)
    arrow(draw, (1020, 330), (1110, 490), TEAL)
    outline_card(draw, (110, 620, 460, 825), "Main Idea", [
        "User actions are handled on the React frontend.",
        "Business logic is processed by Express APIs.",
        "MongoDB stores users, drivers, bookings, and locations.",
    ], NAVY)
    outline_card(draw, (560, 620, 930, 825), "Safety Integration", [
        "Emergency contact details are stored in the user profile.",
        "Ride details can be sent by email or SMS when alerts are enabled.",
    ], ORANGE)
    outline_card(draw, (970, 620, 1490, 825), "Deployment", [
        "Frontend is deployed on Vercel.",
        "Backend is deployed on Railway.",
        "Database is hosted on MongoDB Atlas.",
    ], GREEN)
    img.save(path)


def diagram_workflow(path):
    img, draw = create_canvas("System Workflow", "Ride-booking flow from login to trip completion")
    steps = [
        ("Open App", NAVY),
        ("Login / Signup", TEAL),
        ("Select Pickup & Drop", ORANGE),
        ("Get Fare Estimate", PURPLE),
        ("Create Booking", GREEN),
        ("Assign Driver", NAVY),
        ("Track Ride", TEAL),
        ("Emergency Alert", ORANGE),
        ("Trip Complete", PURPLE),
    ]
    x = 70
    y_positions = [220, 220, 220, 220, 220, 500, 500, 500, 500]
    x_positions = [70, 260, 470, 680, 890, 260, 500, 740, 980]
    sizes = [(160, 90)] * len(steps)
    for (text, color), x, y, (w, h) in zip(steps, x_positions, y_positions, sizes):
        box(draw, (x, y, x + w, y + h), text, color, font=FONT_18_B)
    arrow(draw, (230, 265), (260, 265), TEAL)
    arrow(draw, (420, 265), (470, 265), TEAL)
    arrow(draw, (630, 265), (680, 265), TEAL)
    arrow(draw, (840, 265), (890, 265), TEAL)
    arrow(draw, (970, 310), (340, 500), TEAL)
    arrow(draw, (420, 545), (500, 545), TEAL)
    arrow(draw, (660, 545), (740, 545), TEAL)
    arrow(draw, (900, 545), (980, 545), TEAL)
    outline_card(draw, (1090, 180, 1490, 430), "Workflow Notes", [
        "After booking, the app creates OTP and driver details.",
        "Tracking stages update automatically.",
        "Emergency communication runs if enabled in the profile.",
    ], NAVY)
    img.save(path)


def diagram_lifecycle(path):
    img, draw = create_canvas("Ride Lifecycle", "Status stages used in the tracking module")
    stages = [
        ("Searching", NAVY),
        ("Driver Assigned", TEAL),
        ("Driver Arriving", ORANGE),
        ("Trip in Progress", PURPLE),
        ("Completed", GREEN),
    ]
    x = 80
    for i, (text, color) in enumerate(stages):
        x1 = 80 + i * 295
        box(draw, (x1, 340, x1 + 220, 450), text, color, font=FONT_22_B)
        if i < len(stages) - 1:
            arrow(draw, (x1 + 220, 395), (x1 + 295, 395), TEAL)
    outline_card(draw, (180, 570, 1410, 800), "Explanation", [
        "These stages are used by the RideContext to simulate real trip progress.",
        "Each stage updates the activity log and the tracking display for the rider.",
    ], NAVY)
    img.save(path)


def diagram_dfd(path):
    img, draw = create_canvas("Data Flow Diagram", "Flow of ride data between user, processes, database, and safety services")
    box(draw, (80, 350, 250, 440), "User", NAVY, font=FONT_22_B)
    box(draw, (340, 200, 620, 285), "Authentication Process", TEAL, font=FONT_18_B)
    box(draw, (340, 350, 620, 435), "Ride Booking Process", ORANGE, font=FONT_18_B)
    box(draw, (340, 500, 620, 585), "Tracking Process", PURPLE, font=FONT_18_B)
    box(draw, (750, 170, 1030, 255), "User Collection", GREEN, font=FONT_18_B)
    box(draw, (750, 335, 1030, 420), "Booking Collection", GREEN, font=FONT_18_B)
    box(draw, (750, 500, 1030, 585), "Driver Collection", GREEN, font=FONT_18_B)
    box(draw, (1150, 265, 1480, 360), "Emergency Alert Service", NAVY, font=FONT_18_B)
    box(draw, (1150, 470, 1480, 565), "Emergency Contact", ORANGE, font=FONT_18_B)
    arrow(draw, (250, 390), (340, 240), TEAL)
    arrow(draw, (250, 395), (340, 390), TEAL)
    arrow(draw, (250, 405), (340, 540), TEAL)
    arrow(draw, (620, 240), (750, 210), TEAL)
    arrow(draw, (620, 390), (750, 375), TEAL)
    arrow(draw, (620, 540), (750, 540), TEAL)
    arrow(draw, (620, 390), (1150, 315), TEAL)
    arrow(draw, (1480, 360), (1480, 470), TEAL)
    img.save(path)


def diagram_er(path):
    img, draw = create_canvas("ER Diagram", "Main entities and relationships in the Rydo database")
    outline_card(draw, (60, 180, 360, 610), "USER", [
        "userId",
        "name",
        "email",
        "password",
        "phone",
        "emergencyAlertsEnabled",
        "emergencyContact",
    ], NAVY)
    outline_card(draw, (470, 140, 840, 670), "BOOKING", [
        "bookingId",
        "userId",
        "pickup",
        "dropoff",
        "rideType",
        "payment",
        "tripMode",
        "distance",
        "time",
        "fare",
        "otp",
        "status",
        "activityLog",
    ], TEAL)
    outline_card(draw, (930, 180, 1230, 610), "DRIVER", [
        "driverId",
        "name",
        "vehicle",
        "plate",
        "rating",
        "eta",
        "rideType",
        "status",
    ], ORANGE)
    outline_card(draw, (1280, 180, 1540, 610), "LOCATION", [
        "locationId",
        "name",
        "x",
        "y",
        "order",
    ], PURPLE)
    arrow(draw, (360, 390), (470, 390), TEAL)
    arrow(draw, (840, 390), (930, 390), TEAL)
    arrow(draw, (1230, 390), (1280, 390), TEAL)
    draw.text((390, 350), "1 : M", font=FONT_18_B, fill=NAVY)
    draw.text((865, 350), "M : 1", font=FONT_18_B, fill=NAVY)
    draw.text((1240, 350), "M : 1", font=FONT_18_B, fill=NAVY)
    img.save(path)


def diagram_use_case(path):
    img, draw = create_canvas("Use Case Diagram", "Main interactions between rider, emergency contact, and system")
    draw.rounded_rectangle((220, 145, 1360, 760), radius=26, fill=CARD, outline=LINE, width=3)
    draw.text((690, 160), "RYDO SYSTEM", font=FONT_22_B, fill=NAVY)
    actor(draw, 70, 315, NAVY)
    draw.text((85, 520), "User", font=FONT_18_B, fill=NAVY)
    actor(draw, 1420, 315, ORANGE)
    draw.text((1365, 520), "Emergency Contact", font=FONT_18_B, fill=ORANGE)
    use_cases = [
        ((350, 210, 560, 290), "Register / Login", NAVY),
        ((620, 210, 830, 290), "Manage Profile", TEAL),
        ((900, 210, 1160, 290), "Set Emergency Contact", ORANGE),
        ((350, 380, 560, 460), "Book Ride", TEAL),
        ((620, 380, 860, 460), "View Fare Estimate", PURPLE),
        ((930, 380, 1140, 460), "Track Ride", GREEN),
        ((500, 560, 720, 640), "Receive OTP", NAVY),
        ((850, 560, 1085, 640), "Receive Ride Alert", ORANGE),
    ]
    for coords, text, color in use_cases:
        draw.ellipse(coords, fill=WHITE, outline=color, width=3)
        draw_centered_text(draw, coords, text, FONT_16_B, TEXT)
    for target in [(350, 250), (620, 250), (900, 250), (350, 420), (620, 420), (930, 420), (500, 600)]:
        arrow(draw, (160, 400), target, LINE, 3)
    arrow(draw, (1420, 400), (1085, 600), LINE, 3)
    img.save(path)


def generate_diagrams():
    ASSET_DIR.mkdir(exist_ok=True)
    paths = {
        "### System Architecture Diagram": ASSET_DIR / "system_architecture.png",
        "### Workflow Diagram": ASSET_DIR / "workflow.png",
        "### Ride Lifecycle Diagram": ASSET_DIR / "ride_lifecycle.png",
        "### 3.6.1 Data Flow Diagram": ASSET_DIR / "dfd.png",
        "### 3.6.2 ER Diagram": ASSET_DIR / "er_diagram.png",
        "### 3.6.3 Use Case Diagram": ASSET_DIR / "use_case.png",
    }
    diagram_system_architecture(paths["### System Architecture Diagram"])
    diagram_workflow(paths["### Workflow Diagram"])
    diagram_lifecycle(paths["### Ride Lifecycle Diagram"])
    diagram_dfd(paths["### 3.6.1 Data Flow Diagram"])
    diagram_er(paths["### 3.6.2 ER Diagram"])
    diagram_use_case(paths["### 3.6.3 Use Case Diagram"])
    return paths


def set_default_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for style_name, size in [("Title", 24), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True

    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Right-click and update field to generate the table of contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(text)
    run._r.append(fld_end)


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)


def add_cover_page(document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n\n")
    r = p.add_run(TITLE)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(22)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MINOR PROJECT REPORT")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(18)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n")
    r = p.add_run(DEPARTMENT + "\n" + INSTITUTE)
    r.font.name = "Times New Roman"
    r.font.size = Pt(16)
    r.bold = True

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n\nSubmitted By\n").bold = True
    p.add_run(STUDENT)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\nGuided By\n").bold = True
    p.add_run(MENTOR)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\nAcademic Year: " + ACADEMIC_YEAR)

    document.add_page_break()


def add_certificate_page(document):
    document.add_paragraph("CERTIFICATE", style="Heading 1").alignment = WD_ALIGN_PARAGRAPH.CENTER
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.add_run(
        f"This is to certify that the project report entitled \"{TITLE}\" is a bonafide work carried out by "
        f"{STUDENT} of {DEPARTMENT}, {INSTITUTE}, under the guidance of {MENTOR}. "
        "The work submitted in this report is in partial fulfillment of the requirements for the Minor Project."
    )
    document.add_paragraph("\n\n")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Guide Signature\n\n____________________"
    table.cell(0, 1).text = "Student Signature\n\n____________________"
    table.cell(1, 0).text = "Head of Department\n\n____________________"
    table.cell(1, 1).text = "Date\n\n____________________"
    document.add_page_break()


def add_declaration_page(document):
    document.add_paragraph("DECLARATION", style="Heading 1").alignment = WD_ALIGN_PARAGRAPH.CENTER
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.add_run(
        f"I hereby declare that the project report entitled \"{TITLE}\" submitted by me is an original work "
        "carried out for the Minor Project. The report has not been submitted to any other institution for "
        "the award of any degree or diploma."
    )
    document.add_paragraph("\n\n")
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"{STUDENT}\n(Signature)")
    document.add_page_break()


def add_acknowledgement_page(document):
    document.add_paragraph("ACKNOWLEDGEMENT", style="Heading 1").alignment = WD_ALIGN_PARAGRAPH.CENTER
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.add_run(
        f"I express my sincere gratitude to my mentor, {MENTOR}, for the valuable guidance, encouragement, "
        "and support provided throughout the development of this project. I also thank the faculty members of "
        f"{DEPARTMENT}, {INSTITUTE}, for their suggestions and motivation. I am grateful to my friends and "
        "well-wishers for their support during the completion of this work."
    )
    document.add_page_break()


def extract_abstract(markdown_text):
    match = re.search(r"## Abstract\s+(.*?)\n---", markdown_text, re.S)
    return match.group(1).strip() if match else ""


def add_abstract_page(document, abstract_text):
    document.add_paragraph("ABSTRACT", style="Heading 1").alignment = WD_ALIGN_PARAGRAPH.CENTER
    for chunk in [c.strip() for c in abstract_text.split("\n\n") if c.strip()]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(chunk)
    document.add_page_break()


def add_toc_page(document):
    document.add_paragraph("TABLE OF CONTENTS", style="Heading 1").alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph()
    add_toc(p)
    document.add_page_break()


def parse_markdown_into_doc(document, markdown_text, diagram_paths):
    lines = markdown_text.splitlines()
    start_idx = next(i for i, line in enumerate(lines) if line.strip() == "# CHAPTER 1")
    lines = lines[start_idx:]

    in_mermaid = False
    pending_diagram = None

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if stripped == "```mermaid":
                in_mermaid = True
                continue
            if in_mermaid:
                in_mermaid = False
                if pending_diagram:
                    document.add_picture(str(pending_diagram), width=Inches(6.4))
                    cap = document.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.add_run("Figure: " + Path(pending_diagram).stem.replace("_", " ").title()).italic = True
                    pending_diagram = None
                continue

        if in_mermaid:
            continue

        if stripped == "---":
            document.add_page_break()
            continue

        if not stripped:
            document.add_paragraph("")
            continue

        if stripped.startswith("# "):
            document.add_paragraph(stripped[2:], style="Heading 1")
            continue

        if stripped.startswith("## "):
            document.add_paragraph(stripped[3:], style="Heading 2")
            continue

        if stripped.startswith("### "):
            heading_text = stripped[4:]
            document.add_paragraph(heading_text, style="Heading 3")
            pending_diagram = diagram_paths.get(stripped)
            continue

        if stripped.startswith("- "):
            p = document.add_paragraph(style="List Bullet")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(stripped[2:])
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = document.add_paragraph(style="List Number")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.add_run(re.sub(r"^\d+\.\s", "", stripped))
            continue

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        clean = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped)
        p.add_run(clean)


def build_docx():
    ASSET_DIR.mkdir(exist_ok=True)
    markdown_text = SOURCE_MD.read_text(encoding="utf-8")
    diagram_paths = generate_diagrams()

    document = Document()
    set_default_styles(document)
    add_footer(document.sections[0])

    add_cover_page(document)
    add_certificate_page(document)
    add_declaration_page(document)
    add_acknowledgement_page(document)
    add_abstract_page(document, extract_abstract(markdown_text))
    add_toc_page(document)
    parse_markdown_into_doc(document, markdown_text, diagram_paths)

    for section in document.sections:
        add_footer(section)

    document.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
